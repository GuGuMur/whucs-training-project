"""将实践报告 Markdown 正文写入课程提供的 DOCX 模板。"""

from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parent.parent
TEMPLATE = ROOT / "软件系统实践实验报告模板.docx"
SOURCE = Path(__file__).resolve().parent / "软件系统实践报告正文草稿.md"
OUTPUT = Path(__file__).resolve().parent / "软件系统实践报告.docx"


def set_run_font(run, name: str, size: float, bold: bool = False) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold


def reset_document_body(document: Document) -> None:
    body = document._element.body
    sect_pr = deepcopy(body.sectPr)
    for child in list(body):
        body.remove(child)
    body.append(sect_pr)


def add_field(paragraph, instruction: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "右键更新域"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.extend([begin, instr, separate, text, end])


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.paragraph_format.first_line_indent = Cm(0.74)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size in (("Heading 1", 16), ("Heading 2", 14)):
        style = document.styles[style_name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def add_cover(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Cm(2.5)
    set_run_font(p.add_run("武汉大学"), "黑体", 28, True)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Cm(1.2)
    set_run_font(p.add_run("软件系统实践实验报告"), "黑体", 22, True)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Cm(2.2)
    set_run_font(p.add_run("基于大模型的智能文件管理与智能体协同平台"), "黑体", 18, True)

    document.add_paragraph()
    details = [
        "专业名称：计算机科学与技术",
        "课程名称：〔待填写〕",
        "指导教师：〔待填写〕",
        "学生学号：2025302211124、2025302201015",
        "学生姓名：霍从儒、何因绚",
        "完成日期：2026年7月",
    ]
    for text in details:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(text), "宋体", 14)
    document.add_page_break()


def add_toc(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("目录"), "黑体", 18, True)
    toc = document.add_paragraph()
    toc.paragraph_format.space_before = Pt(12)
    add_field(toc, 'TOC \\o "1-3" \\h \\z \\u')
    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(note.add_run("打开 Word 后请右键目录并选择“更新域”。"), "宋体", 10)
    document.add_page_break()


def add_paragraph_from_markdown(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Normal")
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, "宋体", 12, True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, "宋体", 12)


def add_image_placeholder(document: Document, label: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Cm(15.5)
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F5F7FA")
    tc_pr.append(shading)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tr_height = OxmlElement("w:trHeight")
    tr_height.set(qn("w:val"), "3400")
    tr_height.set(qn("w:hRule"), "atLeast")
    tr_pr.append(tr_height)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Cm(2.2)
    set_run_font(p.add_run("【图片占位】\n" + label), "黑体", 13, True)
    note = cell.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(note.add_run("请在 Word 中替换为对应的实际系统截图"), "宋体", 10)
    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(caption.add_run(label), "宋体", 10)


def add_body(document: Document, markdown: str) -> None:
    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if not line or line.startswith("# ") or line.startswith(">"):
            continue
        if line.startswith("## "):
            p = document.add_paragraph(line[3:], style="Heading 1")
            continue
        if line.startswith("### "):
            p = document.add_paragraph(line[4:], style="Heading 2")
            continue
        if line.startswith("#### "):
            p = document.add_paragraph(line[5:], style="Heading 2")
            continue
        if line.startswith("[[图片占位：") and line.endswith("]]" ):
            add_image_placeholder(document, line[7:-2])
            continue
        if line.startswith("- "):
            p = document.add_paragraph(style="List Paragraph")
            p.style = document.styles["Normal"]
            p.paragraph_format.left_indent = Cm(0.74)
            p.paragraph_format.first_line_indent = Cm(-0.45)
            run = p.add_run("• " + line[2:])
            set_run_font(run, "宋体", 12)
            continue
        add_paragraph_from_markdown(document, line)


def add_page_numbers(document: Document) -> None:
    for section in document.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.clear()
        add_field(p, "PAGE")


def main() -> None:
    document = Document(TEMPLATE)
    reset_document_body(document)
    configure_styles(document)
    add_cover(document)
    add_toc(document)
    add_body(document, SOURCE.read_text(encoding="utf-8"))
    add_page_numbers(document)
    document.settings.element.append(OxmlElement("w:updateFields"))
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
